import base64
import glob
import hashlib
import hmac
import secrets as pysecrets
import html
import io
import json
import os
import threading
import time
import uuid
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo

import gspread
import pandas as pd
import streamlit as st
import extra_streamlit_components as stx
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageOps
from google.oauth2.service_account import Credentials
from gspread.exceptions import WorksheetNotFound

# -------------------------
# 기본 설정
# -------------------------
st.set_page_config(layout="wide", page_title="해링턴 하자 관리", page_icon="✓", initial_sidebar_state="collapsed")

WEB_STATUS_OPTIONS = ["미완료", "확인완료"]
LEGACY_INCOMPLETE_STATUSES = {"미확인", "재확인필요", "미완료", ""}
WEB_STATUS_SHEET_NAME = "web_status"
WEB_ITEMS_SHEET_NAME = "web_items"
WEB_IMAGES_SHEET_NAME = "web_images"
UNIT_ACCOUNTS_SHEET_NAME = "unit_accounts"
# pin_mode가 custom인 경우에만 변경 PIN을 사용합니다. 기존 행(pin_mode 없음)은 초기 PIN 상태로 간주합니다.
MAX_UPLOAD_IMAGES = 5
IMAGE_CHUNK_SIZE = 40000
LEGACY_BUILDING = "204동"
LEGACY_UNIT = "4503호"
MAX_ACTIVE_UNITS = 100
SESSION_IDLE_TIMEOUT_SECONDS = 10 * 60
AUTH_COOKIE_NAME = "apt_check_unit_auth"
AUTH_COOKIE_DAYS = 7
GOOGLE_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

st.markdown(
    """
    <style>
    :root {
        --ink: #0a0a0a;
        --body: #737373;
        --mute: #a3a3a3;
        --canvas: #ffffff;
        --soft: #fafafa;
        --line: #e5e5e5;
        --line-strong: #d4d4d4;
        --dark: #171717;
    }

    html, body, [class*="css"] {
        font-family: ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
        color: var(--ink);
    }
    .stApp { background: var(--canvas); }
    .block-container { max-width: 1180px; padding-top: 3.8rem; padding-bottom: 5rem; }
    header[data-testid="stHeader"] { background: rgba(255,255,255,.96); border-bottom: 1px solid var(--line); }
    [data-testid="stSidebar"] { background: var(--soft); border-right: 1px solid var(--line); }

    h1, h2, h3 { letter-spacing: -0.025em; color: var(--ink); }
    h1 { font-size: 2.15rem !important; font-weight: 600 !important; }
    h2 { font-size: 1.55rem !important; font-weight: 600 !important; }
    h3 { font-size: 1.05rem !important; font-weight: 600 !important; }
    p, .stCaption { color: var(--body); }

    .app-nav {
        display:flex; align-items:center; justify-content:space-between; gap:16px;
        min-height: 48px; padding: 8px 0 24px 0; margin-bottom: 22px; border-bottom:1px solid var(--line);
        overflow: visible;
    }
    .brand-wrap { display:flex; align-items:center; gap:12px; }
    .brand-mark {
        width:34px; height:34px; border-radius:999px; background:var(--ink); color:white;
        display:flex; align-items:center; justify-content:center; font-size:17px; font-weight:700;
    }
    .brand-title { font-size:15px; font-weight:650; color:var(--ink); line-height:1.2; }
    .brand-sub { font-size:12px; color:var(--mute); margin-top:2px; }
    .unit-pill {
        display:inline-flex; align-items:center; gap:7px; padding:8px 14px; border:1px solid var(--line);
        border-radius:999px; background:white; font-size:13px; font-weight:600; color:var(--ink);
    }

    .hero { padding: 28px 0 34px 0; }
    .eyebrow { font-size:13px; color:var(--body); margin-bottom:10px; font-weight:600; }
    .hero-title { font-size:36px; line-height:1.16; font-weight:600; letter-spacing:-0.035em; margin:0; color:var(--ink); }
    .hero-copy { max-width:680px; margin-top:12px; font-size:15px; line-height:1.65; color:var(--body); }

    .login-shell { max-width:680px; margin:8vh auto 0 auto; text-align:center; }
    .login-logo {
        width:54px; height:54px; border-radius:999px; background:var(--ink); color:white;
        display:flex; align-items:center; justify-content:center; margin:0 auto 24px auto; font-size:25px; font-weight:700;
    }
    .login-title { font-size:36px; font-weight:600; line-height:1.12; letter-spacing:-0.04em; margin-bottom:12px; }
    .login-copy { color:var(--body); font-size:15px; margin:0 auto 28px auto; max-width:520px; line-height:1.65; }

    .section-label { font-size:20px; font-weight:600; letter-spacing:-.025em; margin:42px 0 14px 0; }
    .section-caption { color:var(--body); font-size:13px; margin-top:-7px; margin-bottom:16px; }

    .kpi-grid { display:grid; grid-template-columns:repeat(4,minmax(0,1fr)); gap:12px; margin:18px 0 12px 0; }
    .kpi-card { border:1px solid var(--line); border-radius:12px; padding:18px 18px 16px; background:white; }
    .kpi-label { font-size:12px; color:var(--body); margin-bottom:8px; }
    .kpi-value { font-size:27px; line-height:1; font-weight:650; letter-spacing:-.035em; color:var(--ink); }
    .kpi-meta { font-size:11px; color:var(--mute); margin-top:8px; }
    .progress-shell { width:100%; height:8px; background:#f0f0f0; border-radius:999px; overflow:hidden; margin-top:14px; }
    .progress-fill { height:100%; background:var(--ink); border-radius:999px; }

    .issue-card {
        background:white; padding:20px; border-radius:12px; border:1px solid var(--line); margin-bottom:10px; color:var(--ink);
    }
    .issue-head { display:flex; align-items:flex-start; justify-content:space-between; gap:10px; }
    .issue-id { font-size:11px; color:var(--mute); font-family:ui-monospace, SFMono-Regular, Menlo, monospace; margin-bottom:6px; }
    .issue-title { font-size:17px; font-weight:650; letter-spacing:-.02em; margin:0; }
    .issue-detail { font-size:13px; line-height:1.6; color:var(--body); margin:12px 0 0; }
    .status-badge { display:inline-flex; align-items:center; padding:6px 10px; border-radius:999px; font-size:11px; font-weight:650; white-space:nowrap; }
    .status-incomplete { background:#f5f5f5; color:#525252; border:1px solid #e5e5e5; }
    .status-checked { background:#171717; color:#fff; border:1px solid #171717; }
    .status-meta { font-size:11px; color:var(--mute); margin-top:10px; }

    div[data-testid="stForm"] { border:1px solid var(--line) !important; border-radius:12px !important; padding:22px !important; background:white; }
    div[data-testid="stExpander"] {
        border:1.5px solid #b8b8b8 !important;
        border-radius:12px !important;
        background:#f7f7f7 !important;
        box-shadow:none !important;
        overflow:hidden !important;
    }
    div[data-testid="stExpander"] details summary {
        min-height:46px !important;
        padding:0 14px !important;
        font-weight:700 !important;
        color:#111111 !important;
        background:#f3f3f3 !important;
    }
    div[data-testid="stExpander"] details summary:hover {
        background:#ededed !important;
    }
    div[data-testid="stExpander"] details[open] summary {
        background:#eeeeee !important;
        border-bottom:1px solid #d6d6d6 !important;
    }
    div[data-baseweb="select"] > div, .stTextInput input, .stTextArea textarea {
        border-color:var(--line-strong) !important; border-radius:12px !important; box-shadow:none !important;
    }
    .stTextInput input { min-height:42px; }
    .stButton > button, .stDownloadButton > button, .stFormSubmitButton > button {
        border-radius:999px !important; border:1px solid var(--line-strong) !important; background:white !important; color:var(--ink) !important;
        font-weight:600 !important; min-height:40px; box-shadow:none !important;
    }
    .stButton > button[kind="primary"], .stFormSubmitButton > button { background:var(--ink) !important; color:white !important; border-color:var(--ink) !important; }
    .stDownloadButton > button { background:var(--ink) !important; color:white !important; border-color:var(--ink) !important; }
    .stButton > button:disabled { background:var(--soft) !important; color:var(--mute) !important; border-color:var(--line) !important; }
    [data-testid="stFileUploaderDropzone"] { border:1px dashed var(--line-strong) !important; border-radius:12px !important; background:var(--soft) !important; }
    [data-testid="stDataFrame"] { border:1px solid var(--line); border-radius:12px; overflow:hidden; }
    [data-testid="stAlert"] { border-radius:12px !important; box-shadow:none !important; }
    hr { border:none !important; border-top:1px solid var(--line) !important; margin:40px 0 !important; }

    @media (max-width: 760px) {
        .block-container { padding: 3.6rem 1rem 4rem 1rem; }
        .app-nav { padding-bottom:18px; margin-bottom:10px; }
        .hero { padding:20px 0 24px 0; }
        .hero-title, .login-title { font-size:28px; }
        .kpi-grid { grid-template-columns:repeat(2,minmax(0,1fr)); }
        .unit-pill { font-size:12px; padding:7px 11px; }
    }
    </style>
    """,
    unsafe_allow_html=True,
)


# -------------------------
# 세대 로그인 (동/호수) + 동시 활성 세대 제한
# -------------------------
def normalize_building(value):
    value = str(value or "").strip()
    if not value:
        return ""
    return value if value.endswith("동") else f"{value}동"


def normalize_unit(value):
    value = str(value or "").strip()
    if not value:
        return ""
    return value if value.endswith("호") else f"{value}호"


def get_initial_pin(building, unit):
    """세대 초기 PIN: 동 숫자 + 호수 숫자 (예: 204동 4503호 -> 2044503)."""
    building_digits = "".join(ch for ch in str(building or "") if ch.isdigit())
    unit_digits = "".join(ch for ch in str(unit or "") if ch.isdigit())
    if not building_digits or not unit_digits:
        return ""
    return f"{building_digits}{unit_digits}"


# -------------------------
# PIN 인증용 Google Sheet 연결
# -------------------------
def auth_get_secret(name, default=None):
    try:
        return st.secrets.get(name, default)
    except Exception:
        return default


@st.cache_resource
def get_auth_spreadsheet():
    sheet_id = auth_get_secret("WEB_STATUS_SHEET_ID") or os.getenv("WEB_STATUS_SHEET_ID")
    if not sheet_id:
        raise RuntimeError("WEB_STATUS_SHEET_ID가 설정되지 않았습니다.")

    service_account_info = auth_get_secret("gcp_service_account")
    if service_account_info:
        service_account_info = dict(service_account_info)
        if "json_key" in service_account_info:
            raw_json = service_account_info["json_key"]
            if isinstance(raw_json, str):
                service_account_info = json.loads(raw_json)
            elif isinstance(raw_json, dict):
                service_account_info = dict(raw_json)
        credentials = Credentials.from_service_account_info(service_account_info, scopes=GOOGLE_SCOPES)
    elif os.path.exists("credentials.json"):
        credentials = Credentials.from_service_account_file("credentials.json", scopes=GOOGLE_SCOPES)
    else:
        raise RuntimeError("Google 서비스 계정 인증정보가 없습니다.")

    return gspread.authorize(credentials).open_by_key(sheet_id)


def get_unit_accounts_worksheet():
    spreadsheet = get_auth_spreadsheet()
    headers = ["동", "호", "pin_salt", "pin_hash", "updated_at", "pin_mode"]
    try:
        ws = spreadsheet.worksheet(UNIT_ACCOUNTS_SHEET_NAME)
    except WorksheetNotFound:
        ws = spreadsheet.add_worksheet(title=UNIT_ACCOUNTS_SHEET_NAME, rows=1000, cols=len(headers))
        ws.append_row(headers)
    values = ws.get_all_values()
    if not values:
        ws.append_row(headers)
    else:
        current_headers = [str(x).strip() for x in values[0]]
        for h in headers:
            if h not in current_headers:
                current_headers.append(h)
        if current_headers != [str(x).strip() for x in values[0]]:
            ws.update([current_headers], f"A1:{gspread.utils.rowcol_to_a1(1, len(current_headers))}")
    return ws


@st.cache_data(ttl=10)
def load_unit_accounts():
    ws = get_unit_accounts_worksheet()
    records = ws.get_all_records()
    result = {}
    for r in records:
        b = normalize_building(r.get("동", ""))
        u = normalize_unit(r.get("호", ""))
        if not b or not u:
            continue
        result[f"{b}|{u}"] = {
            "building": b,
            "unit": u,
            "pin_salt": str(r.get("pin_salt", "") or "").strip(),
            "pin_hash": str(r.get("pin_hash", "") or "").strip(),
            "pin_mode": str(r.get("pin_mode", "") or "").strip().lower(),
        }
    return result


def _hash_pin(pin, salt_hex):
    salt = bytes.fromhex(salt_hex)
    return hashlib.pbkdf2_hmac("sha256", str(pin).encode("utf-8"), salt, 200_000).hex()


def unit_has_pin(building, unit):
    """사용자가 최초 로그인 후 직접 변경한 PIN이 있는지 확인합니다."""
    rec = load_unit_accounts().get(f"{building}|{unit}")
    return bool(
        rec
        and rec.get("pin_mode") == "custom"
        and rec.get("pin_salt")
        and rec.get("pin_hash")
    )


def unit_pin_mode(building, unit):
    rec = load_unit_accounts().get(f"{building}|{unit}")
    if not rec:
        return "initial"
    mode = str(rec.get("pin_mode", "") or "").strip().lower()
    return mode if mode in {"custom", "reset"} else "initial"


def verify_unit_pin(building, unit, pin):
    """변경 PIN이 있으면 그것을, 없으면 동+호수 초기 PIN을 검증합니다."""
    rec = load_unit_accounts().get(f"{building}|{unit}")
    if rec and rec.get("pin_mode") == "custom" and rec.get("pin_salt") and rec.get("pin_hash"):
        try:
            candidate = _hash_pin(pin, rec["pin_salt"])
            return pysecrets.compare_digest(candidate, rec["pin_hash"])
        except Exception:
            return False

    initial_pin = get_initial_pin(building, unit)
    return bool(initial_pin) and pysecrets.compare_digest(str(pin), initial_pin)


def save_unit_pin(building, unit, pin):
    ws = get_unit_accounts_worksheet()
    salt_hex = pysecrets.token_bytes(16).hex()
    pin_hash = _hash_pin(pin, salt_hex)
    now = datetime.now(ZoneInfo("Asia/Seoul")).strftime("%Y-%m-%d %H:%M:%S")
    values = ws.get_all_values()
    target_row = None
    if values:
        headers = [str(x).strip() for x in values[0]]
        try:
            b_idx = headers.index("동")
            u_idx = headers.index("호")
        except ValueError:
            b_idx, u_idx = 0, 1
        for i, row in enumerate(values[1:], start=2):
            b = normalize_building(row[b_idx] if b_idx < len(row) else "")
            u = normalize_unit(row[u_idx] if u_idx < len(row) else "")
            if b == building and u == unit:
                target_row = i
                break
    row_data = [[building, unit, salt_hex, pin_hash, now, "custom"]]
    if target_row:
        ws.update(row_data, f"A{target_row}:F{target_row}")
    else:
        ws.append_row(row_data[0])
    load_unit_accounts.clear()

def reset_unit_pin(building, unit):
    """관리자용: PIN 인증을 해제해 다음 로그인에 동/호수만으로 진입할 수 있게 합니다."""
    ws = get_unit_accounts_worksheet()
    values = ws.get_all_values()
    if not values:
        return False

    headers = [str(x).strip() for x in values[0]]
    try:
        b_idx = headers.index("동")
        u_idx = headers.index("호")
    except ValueError:
        b_idx, u_idx = 0, 1

    salt_idx = headers.index("pin_salt") if "pin_salt" in headers else 2
    hash_idx = headers.index("pin_hash") if "pin_hash" in headers else 3
    updated_idx = headers.index("updated_at") if "updated_at" in headers else 4
    mode_idx = headers.index("pin_mode") if "pin_mode" in headers else 5

    for row_no, row in enumerate(values[1:], start=2):
        b = normalize_building(row[b_idx] if b_idx < len(row) else "")
        u = normalize_unit(row[u_idx] if u_idx < len(row) else "")
        if b == building and u == unit:
            now = datetime.now(ZoneInfo("Asia/Seoul")).strftime("%Y-%m-%d %H:%M:%S")
            max_col = max(salt_idx, hash_idx, updated_idx, mode_idx) + 1
            padded = list(row) + [""] * max(0, max_col - len(row))
            padded[salt_idx] = ""
            padded[hash_idx] = ""
            padded[updated_idx] = now
            padded[mode_idx] = "reset"
            ws.update([padded[:max_col]], f"A{row_no}:{gspread.utils.rowcol_to_a1(row_no, max_col)}")
            load_unit_accounts.clear()
            return True
    # 계정 행이 없으면 reset 상태의 행을 새로 만듭니다.
    now = datetime.now(ZoneInfo("Asia/Seoul")).strftime("%Y-%m-%d %H:%M:%S")
    ws.append_row([building, unit, "", "", now, "reset"])
    load_unit_accounts.clear()
    return True


def admin_password_ok(password):
    configured = str(auth_get_secret("ADMIN_PASSWORD", "") or "")
    if not configured:
        return False
    return pysecrets.compare_digest(str(password), configured)


def _login_cookie_secret():
    # 별도 비밀키가 있으면 우선 사용하고, 없으면 기존 관리자 비밀번호를 사용합니다.
    return str(auth_get_secret("LOGIN_COOKIE_SECRET", "") or auth_get_secret("ADMIN_PASSWORD", "") or "")


def _unit_auth_fingerprint(building, unit):
    """현재 변경 PIN과 연결된 토큰 지문. PIN 변경/초기화 시 기존 쿠키가 자동 무효화됩니다."""
    rec = load_unit_accounts().get(f"{building}|{unit}")
    if not rec or rec.get("pin_mode") != "custom" or not rec.get("pin_hash"):
        return ""
    raw = f"{building}|{unit}|{rec.get('pin_hash')}|custom"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _make_login_token(building, unit):
    secret = _login_cookie_secret()
    fingerprint = _unit_auth_fingerprint(building, unit)
    if not secret or not fingerprint:
        return ""
    payload = {
        "b": building,
        "u": unit,
        "exp": int(time.time()) + AUTH_COOKIE_DAYS * 24 * 60 * 60,
        "fp": fingerprint,
    }
    raw_json = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    raw = base64.urlsafe_b64encode(raw_json).decode("ascii").rstrip("=")
    sig = hmac.new(secret.encode("utf-8"), raw.encode("ascii"), hashlib.sha256).hexdigest()
    return f"{raw}.{sig}"


def _verify_login_token(token):
    secret = _login_cookie_secret()
    if not secret or not token or "." not in str(token):
        return None
    try:
        raw, sig = str(token).rsplit(".", 1)
        expected = hmac.new(secret.encode("utf-8"), raw.encode("ascii"), hashlib.sha256).hexdigest()
        if not pysecrets.compare_digest(sig, expected):
            return None
        padded = raw + "=" * (-len(raw) % 4)
        payload = json.loads(base64.urlsafe_b64decode(padded.encode("ascii")).decode("utf-8"))
        if int(payload.get("exp", 0)) < int(time.time()):
            return None
        building = normalize_building(payload.get("b", ""))
        unit = normalize_unit(payload.get("u", ""))
        if not building or not unit:
            return None
        current_fp = _unit_auth_fingerprint(building, unit)
        if not current_fp or not pysecrets.compare_digest(str(payload.get("fp", "")), current_fp):
            return None
        return building, unit
    except Exception:
        return None


def get_cookie_manager():
    return stx.CookieManager(key="apt_check_cookie_manager")


def set_login_cookie(cookie_manager, building, unit):
    token = _make_login_token(building, unit)
    if not token:
        return False
    cookie_manager.set(
        AUTH_COOKIE_NAME,
        token,
        expires_at=datetime.now() + timedelta(days=AUTH_COOKIE_DAYS),
        key="set_apt_check_auth",
        path="/",
        same_site="lax",
        secure=True,
    )
    return True


def clear_login_cookie(cookie_manager):
    try:
        cookie_manager.delete(AUTH_COOKIE_NAME, key="delete_apt_check_auth")
    except Exception:
        pass


@st.cache_resource
def get_active_unit_registry():
    # Streamlit 프로세스 내에서 공유되는 경량 활성 세대 레지스트리입니다.
    return {"units": {}, "lock": threading.Lock()}


def _prune_inactive_units(registry, now_ts):
    stale_keys = [
        key
        for key, last_seen in registry["units"].items()
        if now_ts - last_seen > SESSION_IDLE_TIMEOUT_SECONDS
    ]
    for key in stale_keys:
        registry["units"].pop(key, None)


def claim_active_unit(building, unit):
    registry = get_active_unit_registry()
    unit_key = f"{building}|{unit}"
    now_ts = time.time()
    with registry["lock"]:
        _prune_inactive_units(registry, now_ts)
        # 같은 세대가 이미 접속 중이면 새 슬롯을 차지하지 않고 활동시간만 갱신합니다.
        if unit_key in registry["units"]:
            registry["units"][unit_key] = now_ts
            return True, len(registry["units"])
        if len(registry["units"]) >= MAX_ACTIVE_UNITS:
            return False, len(registry["units"])
        registry["units"][unit_key] = now_ts
        return True, len(registry["units"])


def touch_active_unit(building, unit):
    registry = get_active_unit_registry()
    unit_key = f"{building}|{unit}"
    now_ts = time.time()
    with registry["lock"]:
        _prune_inactive_units(registry, now_ts)
        registry["units"][unit_key] = now_ts
        return len(registry["units"])


def release_active_unit(building, unit):
    registry = get_active_unit_registry()
    unit_key = f"{building}|{unit}"
    with registry["lock"]:
        registry["units"].pop(unit_key, None)


COOKIE_MANAGER = get_cookie_manager()

if "logged_in_unit" not in st.session_state:
    st.session_state.logged_in_unit = False

# 모바일 브라우저가 Streamlit 세션을 잃어도 7일 쿠키가 유효하면 자동 로그인합니다.
if not st.session_state.logged_in_unit and not st.session_state.get("pending_login_building"):
    try:
        _saved_token = COOKIE_MANAGER.get(AUTH_COOKIE_NAME)
    except Exception:
        _saved_token = None
    if _saved_token:
        _restored = _verify_login_token(_saved_token)
        if _restored:
            _rb, _ru = _restored
            _admitted, _active_count = claim_active_unit(_rb, _ru)
            if _admitted:
                st.session_state.logged_in_unit = True
                st.session_state.login_building = _rb
                st.session_state.login_unit = _ru
                st.session_state.pin_change_required = False
                st.session_state.cookie_auto_login = True
            else:
                st.session_state.cookie_capacity_blocked = True
        else:
            clear_login_cookie(COOKIE_MANAGER)

if not st.session_state.logged_in_unit:
    pending_building = st.session_state.get("pending_login_building")
    pending_unit = st.session_state.get("pending_login_unit")

    st.markdown(
        """
        <div class="login-shell">
          <div class="login-logo">✓</div>
          <div class="login-title">세대 하자 관리</div>
          <div class="login-copy">동과 호수를 입력해 주세요.</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if pending_building and pending_unit:
        st.markdown(f"**{html.escape(pending_building)} {html.escape(pending_unit)}**")
        try:
            _has_custom_pin = unit_has_pin(pending_building, pending_unit)
        except Exception:
            _has_custom_pin = False
        # 로그인 화면에는 초기/변경 PIN 규칙이나 예시를 노출하지 않습니다.
        with st.form("unit_pin_login_form"):
            login_pin = st.text_input("PIN", type="password", max_chars=12)
            pin_submit = st.form_submit_button("로그인", use_container_width=True, type="primary")
        if st.button("다른 동·호수 입력", use_container_width=True):
            st.session_state.pop("pending_login_building", None)
            st.session_state.pop("pending_login_unit", None)
            st.rerun()
        if pin_submit:
            if not login_pin.isdigit():
                st.error("숫자 PIN을 입력해 주세요.")
            elif not verify_unit_pin(pending_building, pending_unit, login_pin):
                st.error("PIN이 올바르지 않습니다.")
            else:
                try:
                    _logged_in_with_initial_pin = not unit_has_pin(pending_building, pending_unit)
                except Exception:
                    _logged_in_with_initial_pin = False
                admitted, active_count = claim_active_unit(pending_building, pending_unit)
                if not admitted:
                    st.warning(
                        f"현재 접속자가 많습니다. 잠시 후 다시 접속해 주세요. "
                        f"(최대 동시 접속 {MAX_ACTIVE_UNITS}세대)"
                    )
                    st.caption("10분 동안 활동이 없는 세대는 자동으로 접속 자리에서 제외됩니다.")
                else:
                    st.session_state.logged_in_unit = True
                    st.session_state.login_building = pending_building
                    st.session_state.login_unit = pending_unit
                    st.session_state.pin_change_required = _logged_in_with_initial_pin
                    st.session_state.pop("pending_login_building", None)
                    st.session_state.pop("pending_login_unit", None)
                    # 변경 PIN으로 로그인한 경우에만 7일 로그인 쿠키를 발급합니다.
                    if not _logged_in_with_initial_pin:
                        set_login_cookie(COOKIE_MANAGER, pending_building, pending_unit)
                    st.rerun()
    else:
        with st.form("unit_login_form"):
            lc1, lc2 = st.columns(2)
            with lc1:
                login_building_raw = st.text_input("동", placeholder="예: 101 또는 101동")
            with lc2:
                login_unit_raw = st.text_input("호수", placeholder="예: 1111 또는 1111호")
            login_submit = st.form_submit_button("세대 관리 시작", use_container_width=True, type="primary")

        if login_submit:
            login_building = normalize_building(login_building_raw)
            login_unit = normalize_unit(login_unit_raw)
            if not login_building or not login_unit:
                st.error("동과 호수를 모두 입력해 주세요.")
            else:
                try:
                    _mode = unit_pin_mode(login_building, login_unit)
                except Exception:
                    _mode = "initial"
                if _mode == "reset":
                    admitted, active_count = claim_active_unit(login_building, login_unit)
                    if not admitted:
                        st.warning(f"현재 접속자가 많습니다. 잠시 후 다시 접속해 주세요. (최대 동시 접속 {MAX_ACTIVE_UNITS}세대)")
                        st.caption("10분 동안 활동이 없는 세대는 자동으로 접속 자리에서 제외됩니다.")
                    else:
                        st.session_state.logged_in_unit = True
                        st.session_state.login_building = login_building
                        st.session_state.login_unit = login_unit
                        st.session_state.pin_change_required = True
                        st.session_state.pin_reset_login = True
                        st.rerun()
                else:
                    st.session_state.pending_login_building = login_building
                    st.session_state.pending_login_unit = login_unit
                    st.rerun()

    st.markdown("---")
    with st.expander("⚙️ 관리자 PIN 초기화", expanded=False):
        st.caption("PIN을 잊어버린 세대의 PIN 인증을 초기화합니다. 하자 데이터는 삭제되지 않으며, 초기화 후에는 동과 호수만 입력해 로그인할 수 있습니다.")
        with st.form("admin_pin_reset_login_form"):
            ac1, ac2 = st.columns(2)
            with ac1:
                admin_building_raw = st.text_input("초기화할 동", placeholder="예: 101 또는 101동", key="admin_login_building")
            with ac2:
                admin_unit_raw = st.text_input("초기화할 호수", placeholder="예: 1111 또는 1111호", key="admin_login_unit")
            admin_password = st.text_input("관리자 비밀번호", type="password", key="admin_login_password")
            admin_reset_submit = st.form_submit_button("PIN 초기화", use_container_width=True)

        if admin_reset_submit:
            target_b = normalize_building(admin_building_raw)
            target_u = normalize_unit(admin_unit_raw)
            if not auth_get_secret("ADMIN_PASSWORD"):
                st.error("Streamlit Secrets에 ADMIN_PASSWORD가 설정되지 않았습니다.")
            elif not admin_password_ok(admin_password):
                st.error("관리자 비밀번호가 올바르지 않습니다.")
            elif not target_b or not target_u:
                st.error("동과 호수를 모두 입력해 주세요.")
            else:
                try:
                    if reset_unit_pin(target_b, target_u):
                        st.success(f"{target_b} {target_u}의 PIN을 초기화했습니다. 다음 로그인은 PIN 없이 동과 호수만 입력하면 됩니다.")
                    else:
                        st.info("해당 세대에 설정된 PIN 계정을 찾지 못했습니다.")
                except Exception as exc:
                    st.error(f"PIN 초기화에 실패했습니다: {exc}")
    st.stop()

CURRENT_BUILDING = st.session_state.get("login_building", LEGACY_BUILDING)
CURRENT_UNIT = st.session_state.get("login_unit", LEGACY_UNIT)
ACTIVE_UNIT_COUNT = touch_active_unit(CURRENT_BUILDING, CURRENT_UNIT)

# 초기 PIN으로 최초 로그인한 경우, 새 4자리 PIN을 설정해야 관리 화면을 사용할 수 있습니다.
if st.session_state.get("pin_change_required", False):
    _is_reset_login = bool(st.session_state.get("pin_reset_login", False))
    _pin_title = "새 PIN 설정" if _is_reset_login else "초기 PIN 변경"
    _pin_copy = (
        f"{html.escape(CURRENT_BUILDING)} {html.escape(CURRENT_UNIT)}의 PIN이 초기화되었습니다.<br>계속 사용하려면 새 4자리 PIN을 설정해 주세요."
        if _is_reset_login
        else f"{html.escape(CURRENT_BUILDING)} {html.escape(CURRENT_UNIT)}의 최초 로그인입니다.<br>계속 사용하려면 새 4자리 PIN을 설정해 주세요."
    )
    st.markdown(
        f"""
        <div class="login-shell">
          <div class="login-logo">✓</div>
          <div class="login-title">{_pin_title}</div>
          <div class="login-copy">{_pin_copy}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    with st.form("forced_initial_pin_change_form"):
        new_pin1 = st.text_input("새 PIN", type="password", max_chars=4, placeholder="4자리 숫자")
        new_pin2 = st.text_input("새 PIN 확인", type="password", max_chars=4, placeholder="한 번 더 입력")
        forced_change_submit = st.form_submit_button("새 PIN 저장 후 시작", use_container_width=True, type="primary")
    if forced_change_submit:
        if not (new_pin1.isdigit() and len(new_pin1) == 4):
            st.error("새 PIN은 4자리 숫자로 설정해 주세요.")
        elif new_pin1 != new_pin2:
            st.error("PIN 확인 값이 일치하지 않습니다.")
        else:
            try:
                save_unit_pin(CURRENT_BUILDING, CURRENT_UNIT, new_pin1)
                st.session_state.pin_change_required = False
                st.session_state.pop("pin_reset_login", None)
                set_login_cookie(COOKIE_MANAGER, CURRENT_BUILDING, CURRENT_UNIT)
                st.success("새 PIN이 저장되었습니다.")
                st.rerun()
            except Exception as exc:
                st.error(f"PIN 저장에 실패했습니다: {exc}")
    if st.button("로그아웃", key="forced_pin_logout", use_container_width=True):
        clear_login_cookie(COOKIE_MANAGER)
        release_active_unit(CURRENT_BUILDING, CURRENT_UNIT)
        for key in ["logged_in_unit", "login_building", "login_unit", "pending_login_building", "pending_login_unit", "pin_change_required"]:
            st.session_state.pop(key, None)
        st.rerun()
    st.stop()

with st.sidebar:
    st.markdown("### 현재 세대")
    st.markdown(f"**{CURRENT_BUILDING} {CURRENT_UNIT}**")
    st.caption("이 세대의 데이터만 표시됩니다.")

    try:
        _pin_is_set = unit_has_pin(CURRENT_BUILDING, CURRENT_UNIT)
    except Exception:
        _pin_is_set = False

    if not _pin_is_set:
        with st.expander("🔒 PIN 설정", expanded=True):
            st.caption("현재 동+호수 초기 PIN 상태입니다. 새 4자리 PIN으로 변경해 주세요.")
            with st.form("set_unit_pin_form"):
                pin1 = st.text_input("새 PIN", type="password", max_chars=4, placeholder="4자리 숫자")
                pin2 = st.text_input("PIN 확인", type="password", max_chars=4, placeholder="한 번 더 입력")
                set_pin_submit = st.form_submit_button("PIN 설정", use_container_width=True)
            if set_pin_submit:
                if not (pin1.isdigit() and len(pin1) == 4):
                    st.error("PIN은 4자리 숫자로 설정해 주세요.")
                elif pin1 != pin2:
                    st.error("PIN 확인 값이 일치하지 않습니다.")
                else:
                    try:
                        save_unit_pin(CURRENT_BUILDING, CURRENT_UNIT, pin1)
                        set_login_cookie(COOKIE_MANAGER, CURRENT_BUILDING, CURRENT_UNIT)
                        st.success("새 PIN이 설정되었습니다. 다음 로그인부터 변경한 PIN을 입력해 주세요.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"PIN 저장에 실패했습니다: {exc}")
    else:
        with st.expander("🔒 PIN 변경"):
            with st.form("change_unit_pin_form"):
                current_pin = st.text_input("현재 PIN", type="password", max_chars=4)
                new_pin1 = st.text_input("새 PIN", type="password", max_chars=4, placeholder="4자리 숫자")
                new_pin2 = st.text_input("새 PIN 확인", type="password", max_chars=4)
                change_pin_submit = st.form_submit_button("PIN 변경", use_container_width=True)
            if change_pin_submit:
                if not verify_unit_pin(CURRENT_BUILDING, CURRENT_UNIT, current_pin):
                    st.error("현재 PIN이 올바르지 않습니다.")
                elif not (new_pin1.isdigit() and len(new_pin1) == 4):
                    st.error("새 PIN은 4자리 숫자로 설정해 주세요.")
                elif new_pin1 != new_pin2:
                    st.error("새 PIN 확인 값이 일치하지 않습니다.")
                else:
                    try:
                        save_unit_pin(CURRENT_BUILDING, CURRENT_UNIT, new_pin1)
                        set_login_cookie(COOKIE_MANAGER, CURRENT_BUILDING, CURRENT_UNIT)
                        st.success("PIN이 변경되었습니다.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"PIN 변경에 실패했습니다: {exc}")

    with st.expander("⚙️ 관리자"):
        st.caption("세대 PIN을 잊어버린 경우 관리자 권한으로 동+호수 초기 PIN으로 되돌릴 수 있습니다.")
        with st.form("admin_pin_reset_sidebar_form"):
            sb1, sb2 = st.columns(2)
            with sb1:
                sb_building_raw = st.text_input("동", placeholder="101", key="admin_sb_building")
            with sb2:
                sb_unit_raw = st.text_input("호수", placeholder="1111", key="admin_sb_unit")
            sb_admin_password = st.text_input("관리자 비밀번호", type="password", key="admin_sb_password")
            sb_reset_submit = st.form_submit_button("해당 세대 PIN 초기화", use_container_width=True)
        if sb_reset_submit:
            target_b = normalize_building(sb_building_raw)
            target_u = normalize_unit(sb_unit_raw)
            if not auth_get_secret("ADMIN_PASSWORD"):
                st.error("ADMIN_PASSWORD가 설정되지 않았습니다.")
            elif not admin_password_ok(sb_admin_password):
                st.error("관리자 비밀번호가 올바르지 않습니다.")
            elif not target_b or not target_u:
                st.error("동과 호수를 모두 입력해 주세요.")
            else:
                try:
                    if reset_unit_pin(target_b, target_u):
                        if target_b == CURRENT_BUILDING and target_u == CURRENT_UNIT:
                            clear_login_cookie(COOKIE_MANAGER)
                        st.success(f"{target_b} {target_u} PIN을 초기화했습니다.")
                    else:
                        st.info("해당 세대에 설정된 PIN 계정을 찾지 못했습니다.")
                except Exception as exc:
                    st.error(f"PIN 초기화에 실패했습니다: {exc}")

    if st.button("로그아웃", use_container_width=True):
        clear_login_cookie(COOKIE_MANAGER)
        release_active_unit(CURRENT_BUILDING, CURRENT_UNIT)
        for key in ["logged_in_unit", "login_building", "login_unit", "pending_login_building", "pending_login_unit", "pin_change_required"]:
            st.session_state.pop(key, None)
        st.rerun()

# 상단 보조 브랜드 바는 제거하고, 핵심 정보부터 바로 표시합니다.
st.markdown(
    f"""
    <div class="hero">
      <div class="eyebrow">세대 하자 관리</div>
      <div class="hero-title">{html.escape(CURRENT_BUILDING)} {html.escape(CURRENT_UNIT)}</div>
      <div class="hero-copy">하자 등록 · 진행현황 · A/S 신청</div>
    </div>
    """,
    unsafe_allow_html=True,
)

# -------------------------
# 기존 Excel 데이터 로드
# -------------------------
excel_candidates = glob.glob("*.xlsx")
if not excel_candidates:
    st.error("저장소 안에서 엑셀(.xlsx) 파일을 찾지 못했습니다. GitHub에 파일이 실제로 존재하는지 확인해 주세요.")
    st.stop()
EXCEL_PATH = excel_candidates[0]


@st.cache_data
def load_data(mtime, path):
    return pd.read_excel(path, sheet_name="Sheet1")


mtime = os.path.getmtime(EXCEL_PATH)
df = load_data(mtime, EXCEL_PATH)
df.columns = [str(c).strip() for c in df.columns]

# 진행현황 공백 정리 (원본 Excel 파일 자체는 수정하지 않음)
if "진행현황" in df.columns:
    df["진행현황_표시"] = df["진행현황"].fillna("미지정").astype(str).str.strip()
else:
    df["진행현황_표시"] = "미지정"

# 새 하자 등록 목록박스는 원본 Excel의 실제 공간/부위/유형 조합을 기준으로 사용한다.
REFERENCE_DF = df.copy()
for _col in ["공간", "부위", "유형"]:
    if _col not in REFERENCE_DF.columns:
        REFERENCE_DF[_col] = ""
    REFERENCE_DF[_col] = REFERENCE_DF[_col].fillna("").astype(str).str.strip()

REFERENCE_DF = REFERENCE_DF[
    (REFERENCE_DF["공간"] != "") | (REFERENCE_DF["부위"] != "") | (REFERENCE_DF["유형"] != "")
].copy()
REFERENCE_SPACES = sorted([x for x in REFERENCE_DF["공간"].dropna().unique().tolist() if str(x).strip()])
REFERENCE_PARTS = sorted([x for x in REFERENCE_DF["부위"].dropna().unique().tolist() if str(x).strip()])
REFERENCE_TYPES = sorted([x for x in REFERENCE_DF["유형"].dropna().unique().tolist() if str(x).strip()])

# -------------------------
# 웹 전용 상태: Google Sheet
# -------------------------
def get_secret(name, default=None):
    try:
        return st.secrets.get(name, default)
    except Exception:
        return default


@st.cache_resource
def get_google_spreadsheet():
    """Streamlit Secrets/credentials.json으로 Google Spreadsheet 연결."""
    sheet_id = get_secret("WEB_STATUS_SHEET_ID") or os.getenv("WEB_STATUS_SHEET_ID")
    if not sheet_id:
        raise RuntimeError("WEB_STATUS_SHEET_ID가 설정되지 않았습니다.")

    service_account_info = get_secret("gcp_service_account")
    if service_account_info:
        service_account_info = dict(service_account_info)
        if "json_key" in service_account_info:
            raw_json = service_account_info["json_key"]
            if isinstance(raw_json, str):
                service_account_info = json.loads(raw_json)
            elif isinstance(raw_json, dict):
                service_account_info = dict(raw_json)
        credentials = Credentials.from_service_account_info(service_account_info, scopes=GOOGLE_SCOPES)
    elif os.path.exists("credentials.json"):
        credentials = Credentials.from_service_account_file("credentials.json", scopes=GOOGLE_SCOPES)
    else:
        raise RuntimeError("Google 서비스 계정 인증정보가 없습니다.")

    return gspread.authorize(credentials).open_by_key(sheet_id)


def get_or_create_worksheet(name, headers, rows=500, cols=10):
    spreadsheet = get_google_spreadsheet()
    try:
        ws = spreadsheet.worksheet(name)
    except WorksheetNotFound:
        ws = spreadsheet.add_worksheet(title=name, rows=rows, cols=max(cols, len(headers)))
        ws.append_row(headers)
    values = ws.get_all_values()
    if not values:
        ws.append_row(headers)
    else:
        # 기존 시트 스키마를 보존하면서 새 컬럼은 오른쪽에만 추가한다.
        current_headers = [str(x).strip() for x in values[0]]
        for header in headers:
            if header not in current_headers:
                current_headers.append(header)
        if current_headers != [str(x).strip() for x in values[0]]:
            ws.update([current_headers], f"A1:{gspread.utils.rowcol_to_a1(1, len(current_headers))}")
    return ws


@st.cache_resource
def get_web_status_worksheet():
    return get_or_create_worksheet(
        WEB_STATUS_SHEET_NAME,
        ["item_id", "web_status", "updated_at", "item_label"],
        rows=500,
        cols=4,
    )


@st.cache_data(ttl=30)
def load_web_status():
    worksheet = get_web_status_worksheet()
    records = worksheet.get_all_records()
    result = {}
    for record in records:
        item_id = str(record.get("item_id", "")).strip()
        raw_status = str(record.get("web_status", "")).strip()
        # 이전 버전의 `미확인`/`재확인필요` 기록은 삭제하지 않고 모두 `미완료`로 읽는다.
        if raw_status == "확인완료":
            status = "확인완료"
        elif raw_status in LEGACY_INCOMPLETE_STATUSES:
            status = "미완료"
        else:
            status = "미완료"
        if item_id:
            result[item_id] = {
                "status": status,
                "updated_at": str(record.get("updated_at", "")).strip(),
            }
    return result


def save_web_status(item_id, new_status, item_label):
    """item_id 기준으로 진행상태를 insert/update한다. Excel은 건드리지 않는다."""
    if new_status not in WEB_STATUS_OPTIONS:
        raise ValueError(f"지원하지 않는 진행상태입니다: {new_status}")
    worksheet = get_web_status_worksheet()
    item_id = str(item_id).strip()
    now = datetime.now(ZoneInfo("Asia/Seoul")).strftime("%Y-%m-%d %H:%M:%S")

    values = worksheet.get_all_values()
    target_row = None
    for row_num, row in enumerate(values[1:], start=2):
        if row and str(row[0]).strip() == item_id:
            target_row = row_num
            break

    row_data = [[item_id, new_status, now, item_label]]
    if target_row:
        worksheet.update(row_data, f"A{target_row}:D{target_row}")
    else:
        worksheet.append_row(row_data[0])

    load_web_status.clear()


@st.cache_data(ttl=30)
def load_web_items():
    ws = get_or_create_worksheet(
        WEB_ITEMS_SHEET_NAME,
        ["item_id", "공간", "부위", "유형", "상세내용", "created_at", "active", "동", "호"],
        rows=500,
        cols=9,
    )
    records = ws.get_all_records()
    rows = []
    for r in records:
        if str(r.get("active", "TRUE")).strip().upper() in {"FALSE", "0", "N", "NO"}:
            continue
        item_id = str(r.get("item_id", "")).strip()
        if not item_id:
            continue
        building = normalize_building(r.get("동", "")) or LEGACY_BUILDING
        unit = normalize_unit(r.get("호", "")) or LEGACY_UNIT
        rows.append({
            "번호": item_id,
            "공간": str(r.get("공간", "")).strip(),
            "부위": str(r.get("부위", "")).strip(),
            "유형": str(r.get("유형", "")).strip(),
            "상세내용": str(r.get("상세내용", "")).strip(),
            "진행현황_표시": "웹등록",
            "저장된사진파일명": "",
            "데이터출처": "웹등록",
            "등록일시": str(r.get("created_at", "")).strip(),
            "동": building,
            "호": unit,
        })
    return pd.DataFrame(rows)


@st.cache_data(ttl=30)
def load_web_images():
    ws = get_or_create_worksheet(
        WEB_IMAGES_SHEET_NAME,
        ["image_id", "item_id", "sort_order", "chunk_index", "mime_type", "data_chunk"],
        rows=5000,
        cols=6,
    )
    values = ws.get_all_values()
    grouped = {}
    for row in values[1:]:
        if len(row) < 6:
            row = row + [""] * (6 - len(row))
        image_id, item_id, sort_order, chunk_index, mime_type, data_chunk = row[:6]
        if not image_id or not item_id:
            continue
        key = (item_id, image_id)
        g = grouped.setdefault(key, {"sort_order": int(sort_order or 0), "mime_type": mime_type or "image/jpeg", "chunks": []})
        try:
            idx = int(chunk_index or 0)
        except Exception:
            idx = 0
        g["chunks"].append((idx, data_chunk))

    result = {}
    for (item_id, image_id), info in grouped.items():
        encoded = "".join(chunk for _, chunk in sorted(info["chunks"], key=lambda x: x[0]))
        try:
            data = base64.b64decode(encoded)
        except Exception:
            continue
        result.setdefault(item_id, []).append({
            "image_id": image_id,
            "sort_order": info["sort_order"],
            "mime_type": info["mime_type"],
            "data": data,
        })
    for item_id in result:
        result[item_id].sort(key=lambda x: x["sort_order"])
    return result


def compress_image(uploaded_file):
    uploaded_file.seek(0)
    image = Image.open(uploaded_file)
    image = ImageOps.exif_transpose(image).convert("RGB")
    image.thumbnail((1400, 1400))
    out = io.BytesIO()
    image.save(out, format="JPEG", quality=76, optimize=True)
    return out.getvalue(), "image/jpeg"


def save_web_images(item_id, uploaded_files):
    if not uploaded_files:
        return
    ws = get_or_create_worksheet(
        WEB_IMAGES_SHEET_NAME,
        ["image_id", "item_id", "sort_order", "chunk_index", "mime_type", "data_chunk"],
        rows=5000,
        cols=6,
    )
    append_rows = []
    for sort_order, uploaded in enumerate(uploaded_files[:MAX_UPLOAD_IMAGES], start=1):
        data, mime = compress_image(uploaded)
        encoded = base64.b64encode(data).decode("ascii")
        image_id = f"IMG-{uuid.uuid4().hex[:12]}"
        chunks = [encoded[i:i+IMAGE_CHUNK_SIZE] for i in range(0, len(encoded), IMAGE_CHUNK_SIZE)]
        for chunk_index, chunk in enumerate(chunks):
            append_rows.append([image_id, item_id, sort_order, chunk_index, mime, chunk])
    if append_rows:
        ws.append_rows(append_rows, value_input_option="RAW")
    load_web_images.clear()


def create_web_item(space, part, issue_type, detail, initial_status, uploaded_files, building, unit):
    ws = get_or_create_worksheet(
        WEB_ITEMS_SHEET_NAME,
        ["item_id", "공간", "부위", "유형", "상세내용", "created_at", "active", "동", "호"],
        rows=500,
        cols=9,
    )
    now = datetime.now(ZoneInfo("Asia/Seoul")).strftime("%Y-%m-%d %H:%M:%S")
    item_id = f"WEB-{datetime.now(ZoneInfo('Asia/Seoul')).strftime('%Y%m%d%H%M%S')}-{uuid.uuid4().hex[:4].upper()}"
    ws.append_row([item_id, space, part, issue_type, detail, now, "TRUE", building, unit])
    item_label = f"[{item_id}] {space} - {part}"
    save_web_status(item_id, initial_status, item_label)
    save_web_images(item_id, uploaded_files)
    load_web_items.clear()
    load_web_status.clear()
    return item_id


def delete_web_item(item_id):
    """웹에서 새로 등록한 항목만 soft delete(active=FALSE)한다. Excel 원본은 절대 수정하지 않는다."""
    item_id = str(item_id).strip()
    if not item_id.startswith("WEB-"):
        raise ValueError("Excel 원본 항목은 웹에서 삭제할 수 없습니다.")

    ws = get_or_create_worksheet(
        WEB_ITEMS_SHEET_NAME,
        ["item_id", "공간", "부위", "유형", "상세내용", "created_at", "active", "동", "호"],
        rows=500,
        cols=9,
    )
    values = ws.get_all_values()
    target_row = None
    for row_num, row in enumerate(values[1:], start=2):
        if row and str(row[0]).strip() == item_id:
            target_row = row_num
            break

    if not target_row:
        raise ValueError("삭제할 웹 등록 항목을 찾지 못했습니다.")

    # 삭제 이력 복구 가능성을 위해 행/사진은 보존하고 active만 FALSE로 변경한다.
    ws.update([["FALSE"]], f"G{target_row}")
    load_web_items.clear()
    load_web_status.clear()
    load_web_images.clear()


web_status_enabled = True
web_status_error = None
try:
    web_status_map = load_web_status()
except Exception as exc:
    web_status_enabled = False
    web_status_error = str(exc)
    web_status_map = {}

if not web_status_enabled:
    st.warning(
        "진행상태 저장 기능이 아직 연결되지 않았습니다. Google Sheet 연결 설정을 확인해 주세요.\n\n"
        f"설정 내용: {web_status_error}"
    )

# 웹에서 새로 등록한 항목도 함께 불러온다.
try:
    web_items_df = load_web_items() if web_status_enabled else pd.DataFrame()
    web_images_map = load_web_images() if web_status_enabled else {}
except Exception as exc:
    web_items_df = pd.DataFrame()
    web_images_map = {}
    st.warning(f"웹 신규등록 데이터 조회 실패: {type(exc).__name__}: {exc}")

# Excel 원본 데이터는 기존 세대(204동 4503호)의 초기 데이터로 귀속한다.
df["데이터출처"] = "Excel"
df["등록일시"] = ""
df["동"] = LEGACY_BUILDING
df["호"] = LEGACY_UNIT

# 웹 등록 데이터를 병합한다. 기존에 동/호 컬럼 없이 저장된 웹 데이터도 204동 4503호로 자동 귀속된다.
if not web_items_df.empty:
    for col in df.columns:
        if col not in web_items_df.columns:
            web_items_df[col] = ""
    for col in web_items_df.columns:
        if col not in df.columns:
            df[col] = ""
    df = pd.concat([df, web_items_df[df.columns]], ignore_index=True)

# 로그인한 세대의 데이터만 남긴다.
df = df[(df["동"] == CURRENT_BUILDING) & (df["호"] == CURRENT_UNIT)].copy()

# 각 행에 진행상태를 붙인다. Excel 번호 또는 WEB-... ID를 키로 사용한다.
df["진행상태"] = df["번호"].apply(
    lambda x: web_status_map.get(str(x).strip(), {}).get("status", "미완료")
)

# -------------------------
# AS 신청서 생성 (진행상태 기준 / 기존 양식 재현)
# -------------------------
def infer_trade(part_text, type_text=""):
    """기존 AS 신청서의 공종 표기 방식에 맞춰 부위 중심으로 공종을 추정한다."""
    part = str(part_text or "").strip()
    type_value = str(type_text or "").strip()
    combined = f"{part} {type_value}"

    trade_rules = [
        (["벽도배", "천정도배", "도배"], "도배"),
        (["벽도장", "문틀도장", "도장"], "도장"),
        (["마루"], "마루"),
        (["코킹"], "코킹"),
        (["바닥타일", "벽타일", "타일"], "타일"),
        (["스위치", "콘센트", "조명", "전등", "전기"], "전기"),
        (["설비", "배수", "하수구", "수전", "양변기", "세면기"], "설비"),
        (["목문", "목문틀", "문틀", "창문", "창", "신발장", "수납장", "하부장", "상부장", "서랍장", "냉장고장", "가구", "몰딩"], "목공"),
    ]
    for keywords, trade in trade_rules:
        if any(keyword in combined for keyword in keywords):
            return trade
    return "기타"


def _set_cell_text(cell, text, bold=False, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text or ""))
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return p


def _set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")




def _set_table_fixed_widths(table, widths_cm):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid_cols = table._tbl.tblGrid.gridCol_lst
    for idx, width_cm in enumerate(widths_cm):
        width_twips = int(Cm(width_cm).emu / 635)
        if idx < len(grid_cols):
            grid_cols[idx].set(qn("w:w"), str(width_twips))
        if idx < len(table.columns):
            table.columns[idx].width = Cm(width_cm)
        for row in table.rows:
            if idx < len(row.cells):
                _set_cell_width(row.cells[idx], width_cm)

def _prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _collect_item_images(row, item_id):
    images = []
    file_name = str(row.get("저장된사진파일명", "") or "").strip()
    if file_name:
        candidates = [file_name, os.path.basename(file_name)]
        for candidate in candidates:
            if candidate and os.path.exists(candidate):
                try:
                    with open(candidate, "rb") as f:
                        images.append(f.read())
                    break
                except Exception:
                    pass

    for web_img in web_images_map.get(item_id, []):
        data = web_img.get("data")
        if data:
            images.append(data)
    return images


def build_as_request_docx(
    source_df,
    selected_statuses,
    as_date="",
    manager_name="",
    message="",
    building="204동",
    unit="4503호",
    phone="",
):
    """첨부된 기존 A/S 신청서와 유사한 표 형식으로 Word 신청서를 생성한다."""
    selected = source_df[source_df["진행상태"].isin(selected_statuses)].copy()

    doc = Document()
    section = doc.sections[0]
    # 원본 신청서: Letter 용지, 약 1.59cm 여백
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.59)
    section.bottom_margin = Cm(1.59)
    section.left_margin = Cm(1.59)
    section.right_margin = Cm(1.59)

    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("A/S 신청서")
    title_run.bold = True
    title_run.font.size = Pt(18)

    meta_table = doc.add_table(rows=2, cols=3)
    meta_table.style = "Table Grid"
    meta_widths = [3.88, 3.88, 10.65]
    _set_table_fixed_widths(meta_table, meta_widths)
    for idx, label in enumerate(["날짜", "매니저명", "전달사항"]):
        _set_cell_text(meta_table.cell(0, idx), label, bold=True)
    default_msg = f"{', '.join(selected_statuses)} 하자 {len(selected)}건 A/S 요청드립니다."
    _set_cell_text(meta_table.cell(1, 0), as_date or datetime.now(ZoneInfo("Asia/Seoul")).strftime("%m/%d"))
    _set_cell_text(meta_table.cell(1, 1), manager_name)
    _set_cell_text(meta_table.cell(1, 2), message or default_msg)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    info_table = doc.add_table(rows=2, cols=6)
    info_table.style = "Table Grid"
    info_widths = [3.88, 1.86, 1.33, 1.33, 1.33, 8.75]
    labels = ["동", "호", "임시\n키불", "키불출", "입주", "전화"]
    values = [building, unit, "", "", "", phone]
    _set_table_fixed_widths(info_table, info_widths)
    for idx, label in enumerate(labels):
        _set_cell_text(info_table.cell(0, idx), label, bold=True, font_size=8 if idx in [2, 3, 4] else 9)
        _set_cell_text(info_table.cell(1, idx), values[idx])

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    defect_table = doc.add_table(rows=1, cols=4)
    defect_table.style = "Table Grid"
    defect_widths = [1.23, 2.47, 12.24, 2.47]
    _set_table_fixed_widths(defect_table, defect_widths)
    header = defect_table.rows[0]
    _repeat_table_header(header)
    for idx, label in enumerate(["NO", "실", "내용", "공종"]):
        _set_cell_text(header.cells[idx], label, bold=True)

    if selected.empty:
        row_cells = defect_table.add_row().cells
        for c, w in zip(row_cells, defect_widths):
            _set_cell_width(c, w)
        row_cells[0].merge(row_cells[3])
        _set_cell_text(row_cells[0], "선택한 진행상태에 해당하는 하자 항목이 없습니다.")
    else:
        for seq, (_, row) in enumerate(selected.iterrows(), start=1):
            item_id = str(row.get("번호", "") or "").strip()
            space_text = str(row.get("공간", "") or "").strip()
            part_text = str(row.get("부위", "") or "").strip()
            type_text = str(row.get("유형", "") or "").strip()
            detail_text = str(row.get("상세내용", "") or "").strip()
            trade = infer_trade(part_text, type_text)

            defect_row = defect_table.add_row()
            _prevent_row_split(defect_row)
            cells = defect_row.cells
            for c, w in zip(cells, defect_widths):
                _set_cell_width(c, w)
            _set_cell_text(cells[0], seq)
            _set_cell_text(cells[1], space_text)
            _set_cell_text(cells[3], trade)

            cells[2].text = ""
            cells[2].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text_p = cells[2].paragraphs[0]
            text_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            text_p.paragraph_format.space_after = Pt(2)
            text_run = text_p.add_run(f"{space_text} {part_text} / {type_text} - {detail_text}")
            text_run.font.size = Pt(9)

            images = _collect_item_images(row, item_id)
            if images:
                for img_bytes in images:
                    try:
                        img_p = cells[2].add_paragraph()
                        img_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        img_p.paragraph_format.space_before = Pt(1)
                        img_p.paragraph_format.space_after = Pt(1)
                        img_p.add_run().add_picture(io.BytesIO(img_bytes), width=Cm(8.15))
                    except Exception:
                        pass
            else:
                no_img = cells[2].add_paragraph()
                no_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = no_img.add_run("(사진 없음)")
                r.font.size = Pt(8)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue(), len(selected)


# -------------------------
# 진행상태 대시보드
# -------------------------
total = len(df)
done = int((df["진행상태"] == "확인완료").sum())
incomplete = int((df["진행상태"] == "미완료").sum())
progress = int((done / total) * 100) if total > 0 else 0

st.markdown(
    f"""
    <div class="kpi-grid">
      <div class="kpi-card"><div class="kpi-label">총 하자</div><div class="kpi-value">{total}</div><div class="kpi-meta">현재 세대 전체 항목</div></div>
      <div class="kpi-card"><div class="kpi-label">확인완료</div><div class="kpi-value">{done}</div><div class="kpi-meta">처리 확인된 항목</div></div>
      <div class="kpi-card"><div class="kpi-label">미완료</div><div class="kpi-value">{incomplete}</div><div class="kpi-meta">추가 조치가 필요한 항목</div></div>
      <div class="kpi-card"><div class="kpi-label">진행률</div><div class="kpi-value">{progress}%</div><div class="progress-shell"><div class="progress-fill" style="width:{progress}%"></div></div></div>
    </div>
    """,
    unsafe_allow_html=True,
)

web_added_count = int((df["데이터출처"] == "웹등록").sum())
st.caption(f"웹에서 추가한 하자 {web_added_count}건 · 진행상태는 미완료 / 확인완료 두 단계로 관리합니다.")

st.markdown('<div class="section-label">공간별 진행현황</div><div class="section-caption">숫자를 누르면 아래 하자 목록이 해당 조건으로 바로 필터링됩니다.</div>', unsafe_allow_html=True)
if df.empty:
    st.info("표시할 하자 데이터가 없습니다.")
else:
    space_summary = (
        df.groupby("공간", dropna=False)
        .agg(
            전체=("번호", "count"),
            확인완료=("진행상태", lambda x: int((x == "확인완료").sum())),
            미완료=("진행상태", lambda x: int((x == "미완료").sum())),
        )
        .reset_index()
    )
    space_summary["공간"] = space_summary["공간"].fillna("미지정").replace("", "미지정")
    space_summary["진행률값"] = (
        (space_summary["확인완료"] / space_summary["전체"] * 100)
        .fillna(0).round(0).astype(int)
    )
    # 미완료가 많은 공간을 먼저 확인할 수 있도록 미완료 내림차순, 진행률 오름차순으로 정렬
    space_summary = space_summary.sort_values(["미완료", "진행률값", "전체"], ascending=[False, True, False])

    # 하자 목록의 수동 필터와 같은 session_state를 사용해 클릭 필터와 자연스럽게 연동합니다.
    if "issue_space_filter" not in st.session_state:
        st.session_state["issue_space_filter"] = "전체"
    if "issue_progress_filter" not in st.session_state:
        st.session_state["issue_progress_filter"] = "전체"

    h_space, h_total, h_done, h_incomplete, h_rate = st.columns([1.25, 1, 1, 1, 0.9])
    h_space.markdown("**공간**")
    h_total.markdown("**전체**")
    h_done.markdown("**확인완료**")
    h_incomplete.markdown("**미완료**")
    h_rate.markdown("**진행률**")

    for _, summary_row in space_summary.iterrows():
        summary_space = str(summary_row["공간"])
        safe_key = hashlib.sha1(summary_space.encode("utf-8")).hexdigest()[:10]
        c_space, c_total, c_done, c_incomplete, c_rate = st.columns([1.25, 1, 1, 1, 0.9])
        c_space.markdown(f"**{html.escape(summary_space)}**")

        if c_total.button(str(int(summary_row["전체"])), key=f"space_total_{safe_key}", use_container_width=True):
            st.session_state["issue_space_filter"] = summary_space
            st.session_state["issue_progress_filter"] = "전체"
            st.rerun()
        if c_done.button(str(int(summary_row["확인완료"])), key=f"space_done_{safe_key}", use_container_width=True, disabled=int(summary_row["확인완료"]) == 0):
            st.session_state["issue_space_filter"] = summary_space
            st.session_state["issue_progress_filter"] = "확인완료"
            st.rerun()
        if c_incomplete.button(str(int(summary_row["미완료"])), key=f"space_incomplete_{safe_key}", use_container_width=True, disabled=int(summary_row["미완료"]) == 0):
            st.session_state["issue_space_filter"] = summary_space
            st.session_state["issue_progress_filter"] = "미완료"
            st.rerun()
        c_rate.markdown(f"**{int(summary_row['진행률값'])}%**")

    selected_space_now = st.session_state.get("issue_space_filter", "전체")
    selected_progress_now = st.session_state.get("issue_progress_filter", "전체")
    if selected_space_now != "전체" or selected_progress_now != "전체":
        label_parts = []
        if selected_space_now != "전체":
            label_parts.append(selected_space_now)
        if selected_progress_now != "전체":
            label_parts.append(selected_progress_now)
        info_col, clear_col = st.columns([4, 1])
        info_col.info(f"현재 빠른 필터: {' · '.join(label_parts)}")
        if clear_col.button("필터 해제", key="clear_space_summary_filter", use_container_width=True):
            st.session_state["issue_space_filter"] = "전체"
            st.session_state["issue_progress_filter"] = "전체"
            st.rerun()

# -------------------------
# 웹 신규 하자 등록
# -------------------------
with st.expander("새 하자 등록", expanded=False):
    st.caption(
        "새 항목은 현재 로그인한 세대에만 저장되며 Excel을 수정하지 않고 Google Sheet에 별도로 저장됩니다. "
        "공간·부위·유형은 원본 Excel에 있는 전체 목록에서 각각 독립적으로 선택하며 사진은 최대 5장까지 첨부할 수 있습니다."
    )

    if not REFERENCE_SPACES or not REFERENCE_PARTS or not REFERENCE_TYPES:
        st.error("원본 Excel에서 공간/부위/유형 기준값을 읽지 못했습니다.")
        new_space = new_part = new_type = ""
    else:
        c1, c2, c3 = st.columns(3)
        with c1:
            new_space = st.selectbox("공간 *", REFERENCE_SPACES, key="new_issue_space")
        with c2:
            new_part = st.selectbox("부위 *", REFERENCE_PARTS, key="new_issue_part")
        with c3:
            new_type = st.selectbox("유형 *", REFERENCE_TYPES, key="new_issue_type")

    c_status, c_info = st.columns([1, 2])
    with c_status:
        new_initial_status = st.selectbox(
            "등록 후 진행상태",
            WEB_STATUS_OPTIONS,
            index=0,
            key="new_issue_initial_status",
        )
    with c_info:
        st.info(
            f"선택: {new_space or '-'} / {new_part or '-'} / {new_type or '-'}",
            icon="📌",
        )

    new_detail = st.text_area(
        "상세내용 *",
        placeholder="하자 위치와 증상을 자세히 입력해 주세요.",
        height=110,
        key="new_issue_detail",
    )
    new_photos = st.file_uploader(
        "사진 첨부 (최대 5장)",
        type=["jpg", "jpeg", "png", "webp"],
        accept_multiple_files=True,
        help="사진은 Google Sheet에 압축하여 별도 저장됩니다.",
        key="new_issue_photos",
    )

    if new_photos:
        preview_cols = st.columns(min(len(new_photos), 5))
        for p_idx, uploaded in enumerate(new_photos[:MAX_UPLOAD_IMAGES]):
            with preview_cols[p_idx % len(preview_cols)]:
                st.image(uploaded, caption=f"사진 {p_idx + 1}", use_container_width=True)

    submitted = st.button(
        "새 하자 등록",
        use_container_width=True,
        disabled=not web_status_enabled,
        key="submit_new_issue",
        type="primary",
    )

    if submitted:
        if not new_space or not new_part or not new_type or not new_detail.strip():
            st.error("공간, 부위, 유형, 상세내용은 모두 필수입니다.")
        elif len(new_photos or []) > MAX_UPLOAD_IMAGES:
            st.error(f"사진은 최대 {MAX_UPLOAD_IMAGES}장까지 첨부할 수 있습니다.")
        else:
            try:
                new_id = create_web_item(
                    new_space,
                    new_part,
                    new_type,
                    new_detail.strip(),
                    new_initial_status,
                    new_photos or [],
                    CURRENT_BUILDING,
                    CURRENT_UNIT,
                )
                st.success(f"새 하자가 등록되었습니다. ID: {new_id}")
                # 다음 등록을 위해 입력값 일부 초기화
                for _key in ["new_issue_detail", "new_issue_photos"]:
                    if _key in st.session_state:
                        del st.session_state[_key]
                st.rerun()
            except Exception as exc:
                st.error(f"새 하자 등록 실패: {type(exc).__name__}: {exc}")

# AS 신청서 출력: 진행상태 기준
with st.expander("A/S 신청서 출력", expanded=False):
    st.caption(
        "첨부해주신 기존 A/S 신청서처럼 `날짜/매니저명/전달사항`, 세대정보, "
        "`NO/실/내용/공종` 표와 사진이 함께 출력됩니다."
    )

    as_statuses = st.multiselect(
        "출력할 진행상태",
        WEB_STATUS_OPTIONS,
        default=["미완료"],
        key="as_web_statuses",
    )

    c1, c2, c3 = st.columns(3)
    with c1:
        as_date = st.text_input(
            "날짜",
            value=datetime.now(ZoneInfo("Asia/Seoul")).strftime("%m/%d"),
            key="as_date",
        )
        as_building = st.text_input("동", value=CURRENT_BUILDING, disabled=True, key="as_building")
    with c2:
        as_manager = st.text_input("매니저명", value="", key="as_manager")
        as_unit = st.text_input("호", value=CURRENT_UNIT, disabled=True, key="as_unit")
    with c3:
        as_phone = st.text_input("전화", value="", key="as_phone")

    as_target_count = int(df["진행상태"].isin(as_statuses).sum()) if as_statuses else 0
    default_message = f"{', '.join(as_statuses)} 하자 {as_target_count}건 A/S 요청드립니다." if as_statuses else ""
    as_message = st.text_input("전달사항", value=default_message, key="as_message")

    if as_statuses:
        st.info(f"현재 AS 신청서 출력 대상: {as_target_count}건")
        as_docx_bytes, _ = build_as_request_docx(
            df,
            as_statuses,
            as_date=as_date,
            manager_name=as_manager,
            message=as_message,
            building=as_building,
            unit=as_unit,
            phone=as_phone,
        )
        st.download_button(
            "A/S 신청서 다운로드 (.docx)",
            data=as_docx_bytes,
            file_name=f"AS신청서_{datetime.now(ZoneInfo('Asia/Seoul')).strftime('%Y%m%d_%H%M')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    else:
        st.warning("출력할 진행상태를 하나 이상 선택해 주세요.")

st.markdown("---")

st.markdown('<div class="section-label">하자 목록</div><div class="section-caption">공간과 진행상태를 선택해 필요한 항목만 빠르게 확인하세요.</div>', unsafe_allow_html=True)

# -------------------------
# 평면도 / 필터
# -------------------------
with st.expander("공간 위치 확인 (평면도)"):
    if os.path.exists("image_5012c2.jpg"):
        st.image("image_5012c2.jpg", use_container_width=True)
    else:
        st.info("평면도 이미지(image_5012c2.jpg)를 찾을 수 없습니다.")

filter_col1, filter_col2 = st.columns(2)

spaces = sorted([str(x) for x in df["공간"].dropna().unique().tolist()])
space_options = ["전체"] + spaces
progress_options = ["전체"] + WEB_STATUS_OPTIONS

# 데이터 변경으로 기존 선택값이 사라진 경우 안전하게 전체로 되돌립니다.
if st.session_state.get("issue_space_filter", "전체") not in space_options:
    st.session_state["issue_space_filter"] = "전체"
if st.session_state.get("issue_progress_filter", "전체") not in progress_options:
    st.session_state["issue_progress_filter"] = "전체"

with filter_col1:
    space = st.selectbox("공간별 필터링", space_options, key="issue_space_filter")

with filter_col2:
    progress_filter = st.selectbox("진행상태", progress_options, key="issue_progress_filter")

target_df = df.copy()
if space != "전체":
    target_df = target_df[target_df["공간"] == space]
if progress_filter != "전체":
    target_df = target_df[target_df["진행상태"] == progress_filter]

st.caption(f"현재 조건에 {len(target_df)}건이 표시됩니다.")

# -------------------------
# 하자 카드
# -------------------------
cols = st.columns(2)
for i, (index, row) in enumerate(target_df.iterrows()):
    with cols[i % 2]:
        item_id = str(row["번호"]).strip()
        current_web_status = str(row.get("진행상태", "미완료") or "미완료")
        updated_at = web_status_map.get(item_id, {}).get("updated_at", "")

        badge_class = {
            "미완료": "status-incomplete",
            "확인완료": "status-checked",
        }.get(current_web_status, "status-incomplete")
        status_color = {
            "확인완료": "#27ae60",
            "미완료": "#ef4444",
        }.get(current_web_status, "#ef4444")

        space_text = str(row.get("공간", ""))
        part_text = str(row.get("부위", ""))
        type_text = str(row.get("유형", ""))
        detail_text = str(row.get("상세내용", ""))
        item_label = f"[{item_id}] {space_text} - {part_text}"

        safe_id = html.escape(item_id)
        safe_title = html.escape(f"{space_text} · {part_text}")
        safe_type = html.escape(type_text)
        safe_detail = html.escape(detail_text)
        safe_updated = html.escape(str(updated_at))
        st.markdown(
            f"""
            <div class="issue-card">
              <div class="issue-head">
                <div>
                  <div class="issue-id">#{safe_id}</div>
                  <div class="issue-title">{safe_title}</div>
                </div>
                <span class="status-badge {badge_class}">{html.escape(current_web_status)}</span>
              </div>
              <div class="issue-detail"><b>{safe_type}</b> · {safe_detail}</div>
              {f'<div class="status-meta">마지막 변경 · {safe_updated}</div>' if updated_at else ''}
            </div>
            """,
            unsafe_allow_html=True,
        )

        # 버튼을 누르는 즉시 Google Sheet에만 저장한다.
        b1, b2 = st.columns(2)
        with b1:
            if st.button(
                "❌ 미완료",
                key=f"incomplete_{item_id}",
                use_container_width=True,
                disabled=not web_status_enabled or current_web_status == "미완료",
            ):
                try:
                    save_web_status(item_id, "미완료", item_label)
                    st.toast(f"{item_label} → 미완료", icon="❌")
                    st.rerun()
                except Exception as exc:
                    st.error(f"진행상태 저장 실패: {exc}")

        with b2:
            if st.button(
                "✅ 확인완료",
                key=f"checked_{item_id}",
                use_container_width=True,
                disabled=not web_status_enabled or current_web_status == "확인완료",
            ):
                try:
                    save_web_status(item_id, "확인완료", item_label)
                    st.toast(f"{item_label} → 확인완료", icon="✅")
                    st.rerun()
                except Exception as exc:
                    st.error(f"진행상태 저장 실패: {exc}")

        # 삭제는 웹에서 신규 등록한 항목에만 제공한다. Excel 원본 항목은 보호한다.
        if str(row.get("데이터출처", "")) == "웹등록":
            confirm_key = f"delete_confirm_{item_id}"
            if st.session_state.get(confirm_key, False):
                st.warning("이 웹 등록 하자를 목록에서 삭제할까요? 원본 Excel에는 영향이 없습니다.")
                dc1, dc2 = st.columns(2)
                with dc1:
                    if st.button(
                        "🗑️ 정말 삭제",
                        key=f"delete_yes_{item_id}",
                        type="primary",
                        use_container_width=True,
                    ):
                        try:
                            delete_web_item(item_id)
                            st.session_state.pop(confirm_key, None)
                            st.toast(f"{item_label} 삭제 완료", icon="🗑️")
                            st.rerun()
                        except Exception as exc:
                            st.error(f"삭제 실패: {type(exc).__name__}: {exc}")
                with dc2:
                    if st.button(
                        "취소",
                        key=f"delete_no_{item_id}",
                        use_container_width=True,
                    ):
                        st.session_state.pop(confirm_key, None)
                        st.rerun()
            else:
                if st.button(
                    "🗑️ 신규 등록 항목 삭제",
                    key=f"delete_start_{item_id}",
                    use_container_width=True,
                ):
                    st.session_state[confirm_key] = True
                    st.rerun()

        file_name = str(row.get("저장된사진파일명", "")).strip()
        if file_name and os.path.exists(file_name):
            st.image(file_name, use_container_width=True)

        web_imgs = web_images_map.get(item_id, [])
        if web_imgs:
            if len(web_imgs) == 1:
                st.image(web_imgs[0]["data"], use_container_width=True)
            else:
                img_cols = st.columns(min(len(web_imgs), 3))
                for img_idx, web_img in enumerate(web_imgs):
                    with img_cols[img_idx % len(img_cols)]:
                        st.image(web_img["data"], use_container_width=True)